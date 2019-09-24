import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH ## 중앙정렬
from docx.enum.style import WD_STYLE_TYPE ## 스타일
### 단위 ###
from docx.shared import Mm
from docx.shared import Pt
import sys # 문자열 크기
import collections # dict key로 정렬

def get_titles(f,title_cnt):
    titles = []
    dummy = ''
    for i in range(1, title_cnt + 1):
        title = f.readline()
        dummy += title
        while str(i) not in title:
            title = f.readline()
            dummy += title
        titles.append(title.strip().replace(str(i) + '.', '').strip())
    titles = sorted(titles)
    print ('dummy in file : ' + str(sys.getsizeof(dummy)))
    return titles, sys.getsizeof(dummy)

def get_content(f, titles, dummy_byte):
    dummy_byte = dummy_byte + (3 - dummy_byte % 3) # 3의 배수로 올림
    print (dummy_byte)
    f.seek(dummy_byte)
    contents = []

    for title in titles:
        content = ''
        tmp = f.readline()
        while title not in str(tmp):
            if(tmp == ''): # 더 이상 읽을 게 없을 시
                print (title + ' : not found please check')
                break
            tmp = f.readline()
        while True:
            if(tmp == ''): # 더 이상 읽을 게 없을 시
                break
            tmp_content = f.readline()
            tmp_if_title = tmp_content.strip()
            if(tmp_content == ''): # 가장 마지막 시를 읽었을 경우
                break
            if tmp_if_title in titles:
                break
            content += tmp_content
        contents.append(content.strip())
        f.seek(dummy_byte)
    return contents

def get_content_dict(f, titles, dummy_byte):
    dummy_byte = dummy_byte + (3 - dummy_byte % 3) # 3의 배수로 올림
    print (dummy_byte)
    f.seek(dummy_byte)

    titles_content = dict.fromkeys(titles, '')

    for title in list(titles_content.keys()):
        content = ''
        tmp = f.readline()
        while title != str(tmp).strip():
            if(tmp == ''): # 더 이상 읽을 게 없을 시
                print (title + 'not found please check')
                break
            tmp = f.readline()
        while True:
            if(tmp == ''): # 더 이상 읽을 게 없을 시
                print (title + 'not found please check')
                break
            tmp_content = f.readline()
            tmp_if_title = tmp_content.strip()
            if(tmp_content == ''): # 가장 마지막 시를 읽었을 경우
                break
            if tmp_if_title in titles:
                break
            content += tmp_content
        if(titles_content[title] == ''):
            titles_content[title] += content.strip()
        f.seek(dummy_byte)
    return titles_content

def make_docx_list(titles, contents):
    doc = docx.Document()
    ### 문서 기본 설정 ###
    section = doc.sections[0]
    section.page_width = Mm(148)
    section.page_heiht = Mm(225)

    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)

    section.left_margin = Mm(15)
    section.right_margin = Mm(15)

    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    ##################

    ### 문서 스타일 생성 ###
    styles = doc.styles

    title_style = styles.add_style('시 제목',WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = '나눔고딕 ExtraBold'
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕 ExtraBold')
    title_style.font.size = Pt(12)

    content_style = styles.add_style('시 본문',WD_STYLE_TYPE.PARAGRAPH)
    content_style.font.name = '나눔 명조'
    content_style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔 명조')
    content_style.font.size = Pt(11)
    ####################

    for i in range(len(titles)):
        ### 제목 ###
        doc.add_paragraph(titles[i] + '\n', style = title_style).bold = True
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        #### 내용 ###
        doc.add_paragraph(contents[i] + '\n\n', style = content_style)
        
        ### 페이지 넘어가기 ####
        doc.add_page_break()

    return doc

def make_docx_dict(title_content):
    doc = docx.Document()
    ### 문서 기본 설정 ###
    section = doc.sections[0]
    section.page_width = Mm(148)
    section.page_heiht = Mm(225)

    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)

    section.left_margin = Mm(15)
    section.right_margin = Mm(15)

    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    ##################

    ### 문서 스타일 생성 ###
    styles = doc.styles

    title_style = styles.add_style('시 제목',WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = '나눔고딕 ExtraBold'
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕 ExtraBold')
    title_style.font.size = Pt(12)

    content_style = styles.add_style('시 본문',WD_STYLE_TYPE.PARAGRAPH)
    content_style.font.name = '나눔 명조'
    content_style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔 명조')
    content_style.font.size = Pt(11)
    ####################

    for title, content in title_content.items():
        ### 제목 ###
        doc.add_paragraph(title + '\n', style = title_style).bold = True
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #### 내용 ###
        doc.add_paragraph(content + '\n', style = content_style)
        
        ### 페이지 넘어가기 ####
        doc.add_page_break()
    return doc
